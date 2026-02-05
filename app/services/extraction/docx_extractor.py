"""
DOCX Document Extractor

Extracts text from DOCX documents using python-docx library.
Processes both paragraphs and tables to find Forderungshoehe amounts.

Phase 3 Scope:
- Extract text from paragraphs and tables
- Apply same amount patterns as EmailBodyExtractor
- No API calls - pure python-docx extraction
"""

import os
import re
import logging
from typing import List, Dict, Any

from app.models.extraction_result import (
    SourceExtractionResult,
    ExtractedAmount,
    ExtractedEntity,
)
from app.services.extraction.german_preprocessor import GermanTextPreprocessor
from app.services.extraction.german_parser import parse_german_amount
from app.services.extraction.german_validator import GermanValidator

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """
    Extracts structured data from DOCX documents.

    Uses python-docx to extract text from both paragraphs and tables,
    then applies the same regex patterns as EmailBodyExtractor.
    """

    def __init__(self):
        self.preprocessor = GermanTextPreprocessor()
        self.validator = GermanValidator()
        # Amount patterns (same as EmailBodyExtractor)
        self.amount_patterns = [
            # Explicit Gesamtforderung (with flexible separator)
            r'[Gg]esamtforderung[:\s\w]*?([0-9][0-9.,]*)\s*(EUR|€)',
            r'[Gg]esamt(?:betrag|summe)[:\s\w]*?([0-9][0-9.,]*)\s*(EUR|€)',
            # Forderung patterns
            r'[Ff]orderung(?:shöhe|sbetrag)?[:\s\w]*?([0-9][0-9.,]*)\s*(EUR|€)',
            # Generic amount patterns
            r'[Bb]etrag[:\s\w]*?([0-9][0-9.,]*)\s*(EUR|€)',
            r'[Ss]umme[:\s\w]*?([0-9][0-9.,]*)\s*(EUR|€)',
            # Amount followed by currency (catch-all)
            r'([0-9][0-9.,]*)\s*(EUR|€)',
            # Currency first patterns
            r'(EUR|€)\s*([0-9][0-9.,]*)',
        ]

    def extract(self, docx_path: str) -> SourceExtractionResult:
        """
        Extract claim amounts and entity names from DOCX document.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            SourceExtractionResult with extracted amounts and names
        """
        from docx import Document

        result = SourceExtractionResult(
            source_type="docx",
            source_name=os.path.basename(docx_path),
            extraction_method="python_docx",
            tokens_used=0
        )

        try:
            doc = Document(docx_path)

            # Extract all paragraphs
            all_text = []
            paragraph_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
                    paragraph_count += 1

            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = ' | '.join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        all_text.append(row_text)

            combined_text = '\n'.join(all_text)

            # NEW: Apply German preprocessing
            combined_text = self.preprocessor.preprocess(combined_text)

            logger.info(
                f"DOCXExtractor: extracted {paragraph_count} paragraphs, "
                f"{table_count} tables from {os.path.basename(docx_path)}"
            )

            # Extract amounts and names from combined text
            return self._extract_from_text(combined_text, result)

        except Exception as e:
            result.error = f"docx_extraction_error: {str(e)}"
            logger.error(f"DOCXExtractor: failed to extract from {docx_path}: {e}")
            return result

    def _extract_from_text(
        self, text: str, result: SourceExtractionResult
    ) -> SourceExtractionResult:
        """
        Extract amounts and names from combined text.

        Args:
            text: Combined text from paragraphs and tables
            result: SourceExtractionResult to populate

        Returns:
            Updated SourceExtractionResult
        """
        if not text.strip():
            result.error = "empty_docx_content"
            return result

        # Find all amounts
        found_amounts = self._find_amounts(text)

        logger.info(f"DOCXExtractor: found {len(found_amounts)} amounts in document")

        # Take highest amount (USER DECISION: highest wins)
        if found_amounts:
            best = max(found_amounts, key=lambda x: x['value'])
            result.gesamtforderung = ExtractedAmount(
                value=best['value'],
                currency="EUR",
                raw_text=best['raw'],
                source="docx",
                confidence=best['confidence']
            )
            logger.info(
                f"DOCXExtractor: selected {best['value']} EUR "
                f"(confidence: {best['confidence']})"
            )

        # Extract entity names
        self._extract_names(text, result)

        return result

    def _find_amounts(self, text: str) -> List[Dict[str, Any]]:
        """
        Find all monetary amounts in text using regex patterns.

        Returns list of dicts with: value, raw, confidence
        """
        found_amounts = []

        for pattern in self.amount_patterns:
            for match in re.finditer(pattern, text):
                try:
                    # Handle both patterns: "1.234,56 EUR" and "EUR 1.234,56"
                    groups = match.groups()
                    if groups[0] in ('EUR', '€'):
                        amount_str = groups[1]
                    else:
                        amount_str = groups[0]

                    # NEW: Use babel-based parser instead of manual replacement
                    try:
                        amount_value = parse_german_amount(amount_str)
                    except ValueError:
                        continue  # Skip unparseable amounts

                    if amount_value > 0:
                        # Confidence: HIGH if German format detected (has comma decimal)
                        has_german_decimal = ',' in amount_str and amount_str.index(',') > amount_str.rfind('.')
                        found_amounts.append({
                            'value': amount_value,
                            'raw': match.group(0),
                            'confidence': 'HIGH' if has_german_decimal else 'MEDIUM'
                        })
                except (ValueError, IndexError):
                    continue

        return found_amounts

    def _extract_names(self, text: str, result: SourceExtractionResult) -> None:
        """
        Extract client and creditor names from document text.
        """
        name_patterns = [
            (r'(?:Mandant|Schuldner|Kunde)[:\s]+([A-Za-zäöüÄÖÜß\-,\s]+)', 'client'),
            (r'(?:Gläubiger|Inkasso|Firma)[:\s]+([A-Za-zäöüÄÖÜß\-,\s]+)', 'creditor'),
        ]

        for pattern, entity_type in name_patterns:
            match = re.search(pattern, text)
            if match:
                name = match.group(1).strip()
                name = re.sub(r'[,.\s]+$', '', name)

                # NEW: Apply OCR correction to name fields
                name = self.preprocessor.correct_name_field(name)

                # NEW: Validate name format (REQ-GERMAN-04)
                if len(name) > 3 and self.validator.validate_name(name):
                    if entity_type == 'client':
                        result.client_name = ExtractedEntity(
                            value=name,
                            entity_type="client_name",
                            confidence="MEDIUM"
                        )
                    else:
                        result.creditor_name = ExtractedEntity(
                            value=name,
                            entity_type="creditor_name",
                            confidence="MEDIUM"
                        )
                elif len(name) > 3:
                    logger.debug(f"DOCX name '{name}' failed German format validation")
                    if entity_type == 'client':
                        result.client_name = ExtractedEntity(
                            value=name,
                            entity_type="client_name",
                            confidence="LOW"
                        )
                    else:
                        result.creditor_name = ExtractedEntity(
                            value=name,
                            entity_type="creditor_name",
                            confidence="LOW"
                        )

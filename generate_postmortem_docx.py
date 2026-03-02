"""Generate the post-mortem DOCX for the 100 EUR default amount bug."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date


def add_heading_with_style(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def add_key_value(doc, key, value):
    p = doc.add_paragraph()
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    run_key.font.size = Pt(11)
    run_val = p.add_run(str(value))
    run_val.font.size = Pt(11)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def build_document():
    doc = Document()

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # =====================================================================
    # TITLE
    # =====================================================================
    title = doc.add_heading("Post-Mortem: Silent Amount Overwrite Bug", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph("Creditor Email Matcher — Extraction Pipeline")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in subtitle.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # =====================================================================
    # 1. EXECUTIVE SUMMARY
    # =====================================================================
    add_heading_with_style(doc, "1. Executive Summary", level=1)
    doc.add_paragraph(
        "A creditor replied with an email that contained no clear monetary amount. "
        "The extraction pipeline found nothing and silently fell back to a hardcoded "
        "100 EUR default, which then overwrote the correct 430 EUR value already "
        "stored in MongoDB. The client's debt record was corrupted without any "
        "warning or log entry indicating what had happened."
    )

    # =====================================================================
    # 2. IMPACT
    # =====================================================================
    add_heading_with_style(doc, "2. Impact", level=1)

    table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)

    rows_data = [
        ("Severity", "High — silent data corruption"),
        ("Affected Records", "Any email where the creditor reply contained no extractable amount"),
        ("Data Impact", "Existing debt amounts overwritten with 100 EUR default"),
        ("Detection", "Discovered during manual review of a client account"),
        ("User-Facing Effect", "Incorrect debt amounts shown in the Mandanten Portal"),
    ]
    for i, (k, v) in enumerate(rows_data):
        cell_key = table.rows[i].cells[0]
        cell_val = table.rows[i].cells[1]
        cell_key.text = k
        cell_val.text = v
        for paragraph in cell_key.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        for paragraph in cell_val.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    doc.add_paragraph("")

    # =====================================================================
    # 3. ROOT CAUSE ANALYSIS
    # =====================================================================
    add_heading_with_style(doc, "3. Root Cause Analysis", level=1)

    doc.add_paragraph(
        "The bug had two independent root causes that combined to produce the failure:"
    )

    # Root Cause 1
    add_heading_with_style(doc, "Root Cause 1: Hardcoded 100 EUR Default in the Consolidator", level=2)
    doc.add_paragraph(
        "The ExtractionConsolidator class was initialized with a default_amount parameter "
        "set to 100.0. When the extraction pipeline processed all sources (email body and "
        "attachments) and found no monetary amounts in any of them, the consolidator returned "
        "100.0 EUR instead of signaling that no amount was found."
    )

    p = doc.add_paragraph()
    run = p.add_run("Affected code — app/services/extraction/consolidator.py:")
    run.bold = True
    run.font.size = Pt(10)

    code1 = doc.add_paragraph()
    code1.paragraph_format.left_indent = Inches(0.5)
    run = code1.add_run(
        'def __init__(self, default_amount: float = 100.0):\n'
        '    self.default_amount = default_amount\n\n'
        '# ... later, when no amounts found:\n'
        'final_amount = self.default_amount  # Always 100.0'
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)

    doc.add_paragraph(
        "This design decision was originally made as a safe fallback during early development, "
        "but it became dangerous once the pipeline began writing results back to MongoDB — "
        "a fabricated amount is worse than no amount at all."
    )

    # Root Cause 2
    add_heading_with_style(doc, "Root Cause 2: No Guard Before Database Write", level=2)
    doc.add_paragraph(
        "The email processing pipeline wrote the extracted amount directly to MongoDB "
        "without comparing it to the existing value. There was no check for:"
    )
    add_bullet(doc, "Whether the new amount was None (extraction found nothing)")
    add_bullet(doc, "Whether the extraction confidence was high enough to justify a write")
    add_bullet(doc, "Whether the new amount was lower than the existing value (a downgrade)")

    doc.add_paragraph(
        "This meant that any amount returned by the consolidator — including the 100 EUR "
        "default — would silently overwrite whatever was already stored."
    )

    # Secondary contributing factor
    add_heading_with_style(doc, "Contributing Factor: Circuit Breaker Also Used 100 EUR", level=2)
    doc.add_paragraph(
        "The ContentExtractionService had a _make_circuit_breaker_result() method that also "
        "returned gesamtforderung=100.0 when the daily cost circuit breaker was open. "
        "This created a second pathway to the same bug: if the API cost limit was hit, "
        "all subsequent emails would have their amounts overwritten with 100 EUR."
    )

    # =====================================================================
    # 4. TIMELINE OF THE BUG
    # =====================================================================
    add_heading_with_style(doc, "4. How the Bug Manifested", level=1)

    doc.add_paragraph("Step-by-step trace of the original failure:")

    steps = [
        ("1. Email received", "A creditor replied to an inquiry. The email body was a short "
         "acknowledgment with no monetary amount mentioned. No attachments."),
        ("2. Extraction ran", "The pipeline processed the email body. No amounts were found "
         "in any source."),
        ("3. Consolidator defaulted", "ExtractionConsolidator.consolidate() returned "
         "gesamtforderung=100.0 with confidence='LOW'."),
        ("4. No guard existed", "The email processor passed 100.0 directly to the "
         "DualDatabaseWriter without checking the existing value."),
        ("5. MongoDB overwritten", "The existing 430 EUR value was replaced with 100 EUR. "
         "No log entry indicated this was a default, not an extracted value."),
        ("6. Discovery", "A team member noticed the incorrect amount during a manual "
         "account review."),
    ]

    for step_title, step_desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(step_title + " — ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(step_desc)
        run.font.size = Pt(11)

    # =====================================================================
    # 5. FIX IMPLEMENTED
    # =====================================================================
    add_heading_with_style(doc, "5. Fix Implemented", level=1)

    doc.add_paragraph(
        "The fix addresses both root causes and adds defense-in-depth to prevent "
        "similar issues in the future."
    )

    # Fix 1
    add_heading_with_style(doc, "Fix 1: Make Amount Nullable — Return None Instead of 100 EUR", level=2)
    doc.add_paragraph(
        "The ConsolidatedExtractionResult model's gesamtforderung field was changed from "
        "float to Optional[float] = None. The consolidator now returns None when no amounts "
        "are found, instead of fabricating a default. The circuit breaker result was updated "
        "to also return None."
    )
    add_bullet(doc, "app/models/extraction_result.py", bold_prefix="Model: ")
    add_bullet(doc, "app/services/extraction/consolidator.py", bold_prefix="Consolidator: ")
    add_bullet(doc, "app/actors/content_extractor.py", bold_prefix="Circuit breaker: ")

    # Fix 2
    add_heading_with_style(doc, "Fix 2: Amount Update Guard", level=2)
    doc.add_paragraph(
        "A new should_update_amount() guard function was introduced. It sits between the "
        "extraction pipeline and the database write, enforcing three safety rules:"
    )

    guard_table = doc.add_table(rows=5, cols=3, style="Light Grid Accent 1")
    guard_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Rule", "Condition", "Decision"]
    for i, h in enumerate(headers):
        cell = guard_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    guard_rows = [
        ("1", "new_amount is None", "SKIP — extraction found nothing"),
        ("2", "confidence < 0.75", "SKIP — extraction not reliable enough"),
        ("3", "new_amount < existing_amount", "SKIP — would downgrade the record"),
        ("4", "All checks pass", "UPDATE — write approved"),
    ]
    for i, (rule, cond, decision) in enumerate(guard_rows):
        guard_table.rows[i + 1].cells[0].text = rule
        guard_table.rows[i + 1].cells[1].text = cond
        guard_table.rows[i + 1].cells[2].text = decision
        for j in range(3):
            for p in guard_table.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Every guard decision is logged with structured fields (email_id, existing_amount, "
        "new_amount, confidence, decision, reason) for full auditability."
    )
    add_bullet(doc, "app/services/amount_update_guard.py (new file)", bold_prefix="Guard: ")
    add_bullet(doc, "app/actors/email_processor.py", bold_prefix="Integration: ")

    # Fix 3
    add_heading_with_style(doc, "Fix 3: Pass Existing Amounts Through the Pipeline", level=2)
    doc.add_paragraph(
        "The consolidation agent (Agent 3) now reads the existing debt amount from MongoDB "
        "and passes it through as existing_current_debt_amount in the consolidation result. "
        "This allows the email processor to supply it to the guard for comparison."
    )
    add_bullet(doc, "app/actors/consolidation_agent.py", bold_prefix="Changed: ")

    # Fix 4
    add_heading_with_style(doc, "Fix 4: Structured Logging", level=2)
    doc.add_paragraph(
        "The email processor now emits a comprehensive email_processed log event at the end "
        "of every auto-match, including: existing_amount, extracted_amount, "
        "extraction_confidence, update_decision (UPDATED/SKIPPED), skip_reason, and "
        "confidence_route. This makes it trivial to audit any future amount write."
    )

    # =====================================================================
    # 6. TESTING
    # =====================================================================
    add_heading_with_style(doc, "6. Testing", level=1)

    doc.add_paragraph("Two new test suites were added covering all critical scenarios:")

    test_table = doc.add_table(rows=9, cols=3, style="Light Grid Accent 1")
    test_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    test_headers = ["Test ID", "Scenario", "Expected Result"]
    for i, h in enumerate(test_headers):
        cell = test_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    test_rows = [
        ("TC-01", "Extraction returns None", "Guard blocks (extraction_returned_none)"),
        ("TC-02", "Amount downgrade 430 -> 300", "Guard blocks (amount_downgrade_prevented)"),
        ("TC-03", "Low confidence extraction", "Guard blocks (low_extraction_confidence)"),
        ("TC-04", "Multiple amounts in consolidator", "Highest wins, MEDIUM confidence"),
        ("TC-05", "Amount upgrade 430 -> 500", "Guard approves"),
        ("TC-06", "No matching candidates", "Email routed to review queue"),
        ("TC-07", "Forwarded email format", "Content preserved for extraction"),
        ("TC-08", "API timeout (None extraction)", "Guard blocks write"),
    ]
    for i, (tid, scenario, expected) in enumerate(test_rows):
        test_table.rows[i + 1].cells[0].text = tid
        test_table.rows[i + 1].cells[1].text = scenario
        test_table.rows[i + 1].cells[2].text = expected
        for j in range(3):
            for p in test_table.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph("Additionally, 8 consolidator edge-case tests verify:")
    add_bullet(doc, "No sources returns None (not 100 EUR)")
    add_bullet(doc, "All sources with no amounts returns None")
    add_bullet(doc, "Deduplication within 1 EUR tolerance")
    add_bullet(doc, "Original bug scenario (430 EUR preserved when extraction finds nothing)")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Result: 21 new tests pass. Zero regressions in existing test suite.")
    run.bold = True
    run.font.size = Pt(11)

    # =====================================================================
    # 7. PREVENTION
    # =====================================================================
    add_heading_with_style(doc, "7. Lessons Learned & Prevention", level=1)

    lessons = [
        ("Never fabricate data as a default. ",
         "When extraction finds nothing, the correct answer is None — not an invented "
         "number. A fabricated default that looks like real data is the most dangerous kind "
         "of bug because it passes all downstream validations silently."),
        ("Always guard writes against existing state. ",
         "Any pipeline that writes to a database should compare the new value against the "
         "existing value before overwriting. This is especially critical when the write "
         "source (email extraction) is inherently unreliable."),
        ("Log every write decision with full context. ",
         "Structured logs that include both the old and new values, plus the confidence "
         "score and decision reason, make it possible to audit and catch issues before "
         "they compound."),
        ("Defense in depth. ",
         "The fix applies at three layers: the model (nullable type), the consolidator "
         "(no default), and the guard (write check). Any single layer would have prevented "
         "the bug."),
    ]

    for title, body in lessons:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(body)
        run.font.size = Pt(11)

    # =====================================================================
    # 8. FILES CHANGED
    # =====================================================================
    add_heading_with_style(doc, "8. Files Changed", level=1)

    files_table = doc.add_table(rows=9, cols=2, style="Light Grid Accent 1")
    files_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    files_table.columns[0].width = Inches(3.5)
    files_table.columns[1].width = Inches(3.0)

    file_headers = ["File", "Change"]
    for i, h in enumerate(file_headers):
        cell = files_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    file_rows = [
        ("app/models/extraction_result.py", "Made gesamtforderung nullable, added metadata fields"),
        ("app/services/extraction/consolidator.py", "Removed 100 EUR default, returns None"),
        ("app/actors/content_extractor.py", "Circuit breaker returns None instead of 100"),
        ("app/services/amount_update_guard.py", "NEW — Guard function with structured logging"),
        ("app/actors/consolidation_agent.py", "Passes existing_current_debt_amount through"),
        ("app/actors/email_processor.py", "Integrates guard before database write"),
        ("tests/test_amount_update_guard.py", "NEW — 13 guard tests"),
        ("tests/test_extraction_consolidator.py", "NEW — 8 consolidator edge-case tests"),
    ]
    for i, (f, c) in enumerate(file_rows):
        files_table.rows[i + 1].cells[0].text = f
        files_table.rows[i + 1].cells[1].text = c
        for j in range(2):
            for p in files_table.rows[i + 1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # Save
    output_path = "postmortem_amount_overwrite_bug.docx"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
